from __future__ import annotations

import importlib.util
import tempfile
import unittest
from pathlib import Path
from typing import Any

from hypermindlabs.document_parser.adapters.base import build_parse_profile
from hypermindlabs.document_parser.adapters.html import HtmlWebParserAdapter
from hypermindlabs.document_parser.adapters.office import OfficeDocumentParserAdapter
from hypermindlabs.document_parser.adapters.structured_data import StructuredDataParserAdapter
from hypermindlabs.document_parser.adapters.text_markdown import TextMarkdownParserAdapter
from hypermindlabs.document_parser.router import DocumentParserRouter, DocumentParserRoutingError


_HAS_DOCX = bool(importlib.util.find_spec("docx"))
_HAS_OPENPYXL = bool(importlib.util.find_spec("openpyxl"))


def _profile_payload(path: Path, *, file_mime: str = "application/octet-stream") -> dict[str, Any]:
    return {
        "document_source_id": 1001,
        "document_version_id": 1002,
        "storage_object_id": 1003,
        "file_name": path.name,
        "file_mime": file_mime,
        "file_extension": path.suffix,
        "file_path": str(path),
        "file_size_bytes": path.stat().st_size,
        "scope": {
            "owner_member_id": 5,
            "chat_host_id": 55,
            "chat_type": "member",
            "community_id": None,
            "topic_id": None,
            "platform": "web",
        },
    }


class _Config:
    def __init__(self, values: dict[str, Any] | None = None):
        self._values = dict(values or {})

    def runtimeValue(self, path: str, default: Any) -> Any:  # noqa: N802
        return self._values.get(path, default)


class MultiFormatIngestionTests(unittest.TestCase):
    def test_text_markdown_adapter_parses_sections(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "notes.md"
            path.write_text(
                "# Release Notes\n\nThis is a rollout summary.\n\n- item one\n- item two\n",
                encoding="utf-8",
            )
            adapter = TextMarkdownParserAdapter()
            output = adapter.parse(build_parse_profile(_profile_payload(path, file_mime="text/markdown")))

            self.assertEqual(output.get("status"), "parsed")
            element_types = {
                str((section or {}).get("metadata", {}).get("element_type") or "")
                for section in (output.get("sections") or [])
            }
            self.assertIn("heading", element_types)
            self.assertIn("list", element_types)

    def test_html_adapter_strips_script_and_extracts_content(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "page.html"
            path.write_text(
                """
                <html><head><title>Status Dashboard</title><script>ignore_me = true;</script></head>
                <body><h1>Operations</h1><p>Service healthy and running.</p><ul><li>Node A</li></ul></body>
                </html>
                """,
                encoding="utf-8",
            )
            adapter = HtmlWebParserAdapter(config_manager=_Config())
            output = adapter.parse(build_parse_profile(_profile_payload(path, file_mime="text/html")))

            content = str(output.get("content_text") or "")
            self.assertEqual(output.get("status"), "parsed")
            self.assertIn("Operations", content)
            self.assertNotIn("ignore_me", content)

    def test_structured_data_adapter_parses_json_xml_and_csv(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            adapter = StructuredDataParserAdapter(config_manager=_Config())

            json_path = Path(tmp_dir) / "sample.json"
            json_path.write_text('{"team": {"name": "alpha", "size": 4}}', encoding="utf-8")
            json_output = adapter.parse(build_parse_profile(_profile_payload(json_path, file_mime="application/json")))
            self.assertEqual(json_output.get("status"), "parsed")
            self.assertEqual(json_output.get("metadata", {}).get("structured_format"), "json")

            xml_path = Path(tmp_dir) / "sample.xml"
            xml_path.write_text(
                "<root><project name='atlas'><status>active</status></project></root>",
                encoding="utf-8",
            )
            xml_output = adapter.parse(build_parse_profile(_profile_payload(xml_path, file_mime="application/xml")))
            self.assertEqual(xml_output.get("status"), "parsed")
            self.assertEqual(xml_output.get("metadata", {}).get("structured_format"), "xml")

            csv_path = Path(tmp_dir) / "sample.csv"
            csv_path.write_text("id,name,status\n1,alpha,active\n2,beta,pending\n", encoding="utf-8")
            csv_output = adapter.parse(build_parse_profile(_profile_payload(csv_path, file_mime="text/csv")))
            self.assertEqual(csv_output.get("status"), "parsed")
            self.assertEqual(csv_output.get("metadata", {}).get("structured_format"), "csv")

    @unittest.skipUnless(_HAS_DOCX, "python-docx not installed")
    def test_office_adapter_parses_docx(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.docx"
            doc = Document()
            doc.add_heading("Weekly Brief", level=1)
            doc.add_paragraph("The service remained stable.")
            doc.save(str(path))

            adapter = OfficeDocumentParserAdapter(config_manager=_Config())
            output = adapter.parse(
                build_parse_profile(
                    _profile_payload(
                        path,
                        file_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                )
            )
            self.assertEqual(output.get("status"), "parsed")
            self.assertEqual(output.get("metadata", {}).get("office_format"), "docx")

    @unittest.skipUnless(_HAS_OPENPYXL, "openpyxl not installed")
    def test_office_adapter_parses_xlsx(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "report.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Summary"
            sheet.append(["name", "status"])
            sheet.append(["alpha", "active"])
            workbook.save(str(path))

            adapter = OfficeDocumentParserAdapter(config_manager=_Config())
            output = adapter.parse(
                build_parse_profile(
                    _profile_payload(
                        path,
                        file_mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                )
            )
            self.assertEqual(output.get("status"), "parsed")
            self.assertEqual(output.get("metadata", {}).get("office_format"), "xlsx")

    def test_office_adapter_returns_actionable_errors_for_invalid_pptx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "slides.pptx"
            path.write_bytes(b"not-a-valid-pptx")

            adapter = OfficeDocumentParserAdapter(config_manager=_Config())
            output = adapter.parse(
                build_parse_profile(
                    _profile_payload(
                        path,
                        file_mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                    )
                )
            )
            self.assertEqual(output.get("status"), "failed")
            errors = [str(item) for item in output.get("errors") or []]
            self.assertTrue(any(error.startswith("office_parse_error:") for error in errors))

    def test_router_selects_new_multiformat_adapters(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            markdown_path = Path(tmp_dir) / "notes.md"
            markdown_path.write_text("# Plan\n\n- draft\n", encoding="utf-8")
            html_path = Path(tmp_dir) / "page.html"
            html_path.write_text("<html><body><h1>Team</h1><p>Status green</p></body></html>", encoding="utf-8")
            json_path = Path(tmp_dir) / "state.json"
            json_path.write_text('{"ok": true, "count": 2}', encoding="utf-8")

            router = DocumentParserRouter(
                config_manager=_Config(
                    {
                        "documents.parser_preferred_adapters": [
                            "text-markdown",
                            "html-web",
                            "structured-data",
                            "text-plain",
                            "binary-fallback",
                        ],
                        "documents.parser_min_confidence": 0.2,
                    }
                )
            )

            markdown_result = router.parse_document(_profile_payload(markdown_path, file_mime="text/markdown"))
            html_result = router.parse_document(_profile_payload(html_path, file_mime="text/html"))
            json_path_payload = _profile_payload(json_path, file_mime="application/json")
            json_result = router.parse_document(json_path_payload)

            self.assertEqual(markdown_result.get("provenance", {}).get("selected_adapter"), "text-markdown")
            self.assertEqual(html_result.get("provenance", {}).get("selected_adapter"), "html-web")
            self.assertEqual(json_result.get("provenance", {}).get("selected_adapter"), "structured-data")
            self.assertEqual(json_path_payload["file_extension"], ".json")

    def test_router_enforces_format_allowlist_policy(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "notes.md"
            path.write_text("# Notes\n\nBlocked", encoding="utf-8")
            router = DocumentParserRouter(
                config_manager=_Config(
                    {
                        "documents.parser_format_allowlist": ["pdf"],
                    }
                )
            )
            with self.assertRaises(DocumentParserRoutingError) as ctx:
                router.parse_document(_profile_payload(path, file_mime="text/markdown"))
            self.assertIn("allowlist_miss", str(ctx.exception))


if __name__ == "__main__":
    unittest.main()
